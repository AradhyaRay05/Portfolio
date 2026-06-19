import docx, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

files = [
    r'D:\Education\Coding\Web Dev Projects\My Portfolio\All Resume\Aradhya Ray FWD Resume.docx',
]

for f in files:
    print(f"\n{'='*80}")
    print(f"FILE: {f}")
    print('='*80)
    try:
        doc = docx.Document(f)
        for p in doc.paragraphs:
            if p.text.strip():
                print(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    print(" | ".join(cells))
    except Exception as e:
        print(f"Error: {e}")
